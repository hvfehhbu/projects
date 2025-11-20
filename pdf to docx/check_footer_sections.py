from docx import Document

doc = Document('lab2.docx')

print(f"Number of sections: {len(doc.sections)}")
print("\nChecking footer content in each section:")

for i, section in enumerate(doc.sections):
    footer = section.footer
    print(f"\nSection {i} footer paragraphs:")
    for j, p in enumerate(footer.paragraphs):
        if p.text.strip():
            print(f"  Paragraph {j}: '{p.text}'")
    
    if not any(p.text.strip() for p in footer.paragraphs):
        print(f"  (No text in footer)")
