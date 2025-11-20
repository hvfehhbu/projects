import re
import os
from docx import Document
from docx.shared import Pt

def is_heading(line, patterns):
    """Check if a line matches any of the heading patterns."""
    return any(re.match(pattern, line.strip()) for pattern in patterns)

def create_word_doc(input_filename="ocr_output.txt"):
    try:
        with open(input_filename, 'r', encoding='utf-8') as f:
            ocr_text = f.read()
    except FileNotFoundError:
        print(f"Error: Input file '{input_filename}' not found.")
        return

    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Define heading patterns
    h1_patterns = [r"^CHƯƠNG \w+"]
    h2_patterns = [r"^[IVX]+\..+"] # Matches I., II., III., etc.
    h3_patterns = [r"^\d+\..+"] # Matches 1., 2., 3., etc.
    
    # Clean and split the text into lines
    lines = [line.strip() for line in ocr_text.strip().split('\n') if not line.strip().startswith("==")]

    # Process lines
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line:
            i += 1
            continue

        # Simple heuristic for tables on pages 34-36
        if "CHỦ THỂ BAN HÀNH" in line and "Văn bản QPPL" in line:
            doc.add_paragraph("Bảng: Chủ thể ban hành và văn bản QPPL", style='Heading 3')
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'STT'
            hdr_cells[1].text = 'CHỦ THỂ BAN HÀNH'
            hdr_cells[2].text = 'Văn bản QPPL'
            
            # Manually add table data based on the known structure
            table_data = [
                ["1", "Quốc hội", "1. Hiến pháp; 2. Bộ luật; 3. Luật; 4. Nghị quyết"],
                ["2", "Ủy ban thường vụ Quốc hội", "1. Pháp lệnh; 2. Nghị quyết; 3. Nghị quyết liên tịch"],
                ["3", "Chủ tịch nước", "Lệnh; Quyết định"],
                ["4", "Chính phủ", "Nghị định; Nghị quyết; Nghị quyết liên tịch"],
                ["5", "Thủ tướng Chính phủ", "Quyết định"],
                ["6", "Hội đồng thẩm phán TAND tối cao", "Nghị quyết"],
                ["7", "Chánh án TAND tối cao, Viện trưởng VKSND tối cao", "Thông tư; Thông tư liên tịch"],
                ["8", "Bộ trưởng, Thủ trưởng cơ quan ngang Bộ", "Thông tư; Thông tư liên tịch"],
                ["9", "Tổng kiểm toán Nhà Nước", "Quyết định"],
                ["10", "HĐND cấp tỉnh", "Nghị quyết"],
                ["11", "UBND cấp tỉnh", "Quyết định"],
                ["12", "Đoàn Chủ tịch Ủy ban Trung ương Mặt trận Tổ quốc Việt Nam", "Nghị quyết liên tịch"]
            ]

            for item in table_data:
                row_cells = table.add_row().cells
                row_cells[0].text = item[0]
                row_cells[1].text = item[1]
                row_cells[2].text = item[2]
            
            # Skip the lines that were parsed as a table
            i += 30 # Approximate number of lines for the tables in the OCR text
            doc.add_page_break()
            continue

        # Check for heading types
        if is_heading(line, h1_patterns):
            doc.add_heading(line, level=1)
        elif is_heading(line, h2_patterns):
            doc.add_heading(line, level=2)
        elif is_heading(line, h3_patterns):
            doc.add_heading(line, level=3)
        # Specific cases for titles without clear patterns
        elif line in ["VÍ DỤ GIẢ ĐỊNH", "VÍ DỤ QUY ĐỊNH", "XÁC ĐỊNH BỘ PHẬN QUY ĐỊNH?", "VÍ DỤ CHẾ TÀI", "LƯU Ý", "Những loại chế tài"]:
            doc.add_heading(line, level=4)
        else:
            # This is a paragraph
            paragraph_text = line
            # Combine subsequent lines that don't look like headings or list items
            while (i + 1 < len(lines) and lines[i+1] and 
                   not is_heading(lines[i+1], h1_patterns + h2_patterns + h3_patterns) and
                   not re.match(r"^\d+\)", lines[i+1]) and # e.g. 1)
                   not re.match(r"^[a-z]\)", lines[i+1]) and # e.g. a)
                   not lines[i+1].startswith('-')):
                paragraph_text += " " + lines[i+1].strip()
                i += 1
            doc.add_paragraph(paragraph_text)
        
        i += 1

    doc.save("bao_cao.docx")
    print("Word document 'bao_cao.docx' created successfully.")

if __name__ == "__main__":
    create_word_doc()
    # Clean up the temporary file
    if os.path.exists("ocr_output.txt"):
        os.remove("ocr_output.txt")
        print("Temporary file 'ocr_output.txt' removed.")
